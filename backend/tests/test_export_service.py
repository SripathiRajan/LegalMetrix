import io
import pytest
from datetime import datetime, timezone
import openpyxl
from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from fastapi.testclient import TestClient

from app.db.base import Base
from app.db.models import ScanRecord, Officer
from app.db.scan_repository import ScanRepository
from app.db.session import get_db
from app.auth.dependencies import get_current_active_officer
from app.services.export_service import ExportService
from app.main import app

client = TestClient(app)


@pytest.fixture
def in_memory_db():
    """Sets up an isolated in-memory SQLite database session for testing."""
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
        echo=False
    )
    Base.metadata.create_all(bind=engine)
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    session = TestingSessionLocal()
    try:
        yield session
    finally:
        session.close()
        Base.metadata.drop_all(bind=engine)


@pytest.fixture
def sample_scan_data():
    """Provides representative scan assessment data for export testing."""
    return {
        "id": 105,
        "product_name": "Premium Basmati Rice 1kg",
        "overall_status": "NON_COMPLIANT",
        "compliance_score": 75.0,
        "created_at": datetime(2026, 8, 20, 10, 15, 0, tzinfo=timezone.utc),
        "officer_id": 42,
        "compliance_result": {
            "overall_status": "NON_COMPLIANT",
            "compliance_score": 75.0,
            "total_checks": 4,
            "passed": 3,
            "failed": 1,
            "warnings": 0,
            "results": [
                {
                    "rule_id": "LMPC_RULE_6_1_E",
                    "declaration": "Maximum Retail Price (MRP)",
                    "status": "FAIL",
                    "detected_value": "₹120.00",
                    "reason": "Missing mandatory 'inclusive of all taxes' declaration",
                    "legal_reference": "Rule 6(1)(e) of LMPC Rules, 2011",
                    "severity": "HIGH"
                },
                {
                    "rule_id": "LMPC_RULE_6_1_C",
                    "declaration": "Net Quantity",
                    "status": "PASS",
                    "detected_value": "1 kg",
                    "reason": "Standard SI metric unit declared",
                    "legal_reference": "Rule 6(1)(c) of LMPC Rules, 2011",
                    "severity": "HIGH"
                },
                {
                    "rule_id": "LMPC_RULE_6_1_A",
                    "declaration": "Manufacturer Name & Address",
                    "status": "PASS",
                    "detected_value": "Heritage Agri Foods Ltd, Karnal, Haryana",
                    "reason": "Complete name and address found",
                    "legal_reference": "Rule 6(1)(a) of LMPC Rules, 2011",
                    "severity": "HIGH"
                },
                {
                    "rule_id": "LMPC_RULE_6_1_D",
                    "declaration": "Date of Manufacture",
                    "status": "PASS",
                    "detected_value": "07/2026",
                    "reason": "Valid month and year declared",
                    "legal_reference": "Rule 6(1)(d) of LMPC Rules, 2011",
                    "severity": "MEDIUM"
                }
            ]
        },
        "extracted_data": {
            "product_name": {"value": "Premium Basmati Rice 1kg", "confidence": 0.98},
            "manufacturer_name": {"value": "Heritage Agri Foods Ltd", "confidence": 0.95},
            "manufacturer_address": {"value": "Plot 14, Industrial Area, Karnal, Haryana 132001", "confidence": 0.92},
            "net_quantity": {"value": "1 kg", "confidence": 0.99},
            "mrp": {"value": "₹120.00", "confidence": 0.97},
            "date_declaration": {"value": "07/2026", "confidence": 0.94},
            "country_of_origin": {"value": "India", "confidence": 0.99}
        },
        "visual_statistics": {
            "pdp_area_ratio": 0.42,
            "mean_font_size_px": 15.5,
            "contrast_ratio": 7.2,
            "readability_score": 94.0,
            "bounding_box_count": 6
        },
        "authenticity_result": {
            "verdict": "GENUINE_LIKELY",
            "similarity_score": 0.96,
            "brand_name": "Heritage Agri Foods",
            "notes": "Packaging graphics and registered trademark match official repository."
        }
    }


# 1. Test CSV Generation
def test_export_service_csv(sample_scan_data):
    service = ExportService()
    csv_text = service.generate_scan_csv(sample_scan_data)

    assert isinstance(csv_text, str)
    assert len(csv_text) > 200

    # Verify section titles
    assert "=== LEGAL METROLOGY COMPLIANCE AUDIT RECORD ===" in csv_text
    assert "=== EXTRACTED MANDATORY DECLARATIONS ===" in csv_text
    assert "=== RULE-BY-RULE STATUTORY ASSESSMENT ===" in csv_text

    # Verify metadata and rule contents
    assert "Premium Basmati Rice 1kg" in csv_text
    assert "NON_COMPLIANT" in csv_text
    assert "75.0%" in csv_text
    assert "Heritage Agri Foods Ltd" in csv_text
    assert "LMPC_RULE_6_1_E" in csv_text
    assert "Missing mandatory 'inclusive of all taxes'" in csv_text


# 2. Test Single Scan Excel (.xlsx) Generation
def test_export_service_xlsx_single(sample_scan_data):
    service = ExportService()
    xlsx_bytes = service.generate_scan_xlsx(sample_scan_data)

    assert isinstance(xlsx_bytes, bytes)
    assert len(xlsx_bytes) > 2000
    # ZIP magic bytes for OpenXML documents
    assert xlsx_bytes.startswith(b"PK\x03\x04")

    # Parse and inspect workbook structure
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
    sheet_names = wb.sheetnames
    assert "Summary" in sheet_names
    assert "Extracted Fields" in sheet_names
    assert "Rule Results" in sheet_names
    assert "Visual Statistics" in sheet_names

    # Check Summary sheet contents
    ws_summary = wb["Summary"]
    assert "LEGAL METROLOGY COMPLIANCE AUDIT REPORT" in str(ws_summary["A1"].value)
    
    # Check Extracted Fields sheet
    ws_fields = wb["Extracted Fields"]
    assert ws_fields.max_row >= 5
    field_names = [ws_fields.cell(row=r, column=2).value for r in range(2, ws_fields.max_row + 1)]
    assert any("Manufacturer" in str(fn) for fn in field_names)

    # Check Rule Results sheet
    ws_rules = wb["Rule Results"]
    assert ws_rules.max_row >= 4
    rule_ids = [ws_rules.cell(row=r, column=1).value for r in range(2, ws_rules.max_row + 1)]
    assert "LMPC_RULE_6_1_E" in rule_ids


# 3. Test Bulk Scans Excel (.xlsx) Generation
def test_export_service_xlsx_bulk(sample_scan_data):
    service = ExportService()
    
    # Create multiple scan records
    scan2 = dict(sample_scan_data)
    scan2["id"] = 106
    scan2["product_name"] = "Organic Green Tea 250g"
    scan2["overall_status"] = "COMPLIANT"
    scan2["compliance_score"] = 100.0

    bulk_bytes = service.generate_bulk_scans_xlsx([sample_scan_data, scan2])

    assert isinstance(bulk_bytes, bytes)
    assert len(bulk_bytes) > 2000
    assert bulk_bytes.startswith(b"PK\x03\x04")

    wb = openpyxl.load_workbook(io.BytesIO(bulk_bytes))
    assert "All Scans" in wb.sheetnames
    assert "Violations Log" in wb.sheetnames

    ws_all = wb["All Scans"]
    assert ws_all.max_row == 3  # Header + 2 scans

    ws_viol = wb["Violations Log"]
    assert ws_viol.max_row >= 2  # Header + at least 1 violation from scan 1


# 4. Test DOCX Show-Cause Notice Draft Generation
def test_export_service_docx_show_cause(sample_scan_data):
    service = ExportService()
    docx_bytes = service.generate_show_cause_docx(sample_scan_data, officer_name="Inspector S. Verma")

    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 2000
    # ZIP magic bytes for OpenXML .docx
    assert docx_bytes.startswith(b"PK\x03\x04")

    # Load and inspect DOCX document
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join([p.text for p in doc.paragraphs])

    assert "GOVERNMENT OF INDIA" in full_text
    assert "DEPARTMENT OF CONSUMER AFFAIRS" in full_text
    assert "SHOW-CAUSE NOTICE UNDER SECTION 15 & 36" in full_text
    assert "Premium Basmati Rice 1kg" in full_text
    assert "Heritage Agri Foods Ltd" in full_text
    assert "Inspector S. Verma" in full_text
    assert "FIFTEEN (15) DAYS" in full_text.upper()

    # Verify table of contraventions exists
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    table_text = " ".join([cell.text for row in table.rows for cell in row.cells])
    assert "LMPC_RULE_6_1_E" in table_text
    assert "Maximum Retail Price (MRP)" in table_text


# 5. Test API Export Endpoints
def test_api_export_endpoints(in_memory_db, sample_scan_data):
    def override_get_db():
        try:
            yield in_memory_db
        finally:
            pass

    app.dependency_overrides[get_db] = override_get_db
    app.dependency_overrides[get_current_active_officer] = lambda: Officer(
        id=1, username="test_inspector", role="inspector", is_active=True
    )

    # Persist sample scan in in-memory test DB
    record = ScanRepository.save_scan(
        db=in_memory_db,
        compliance_result=sample_scan_data["compliance_result"],
        extracted_data=sample_scan_data["extracted_data"],
        authenticity_result=sample_scan_data["authenticity_result"],
        visual_statistics=sample_scan_data["visual_statistics"],
        product_name=sample_scan_data["product_name"],
        overall_status=sample_scan_data["overall_status"],
        compliance_score=sample_scan_data["compliance_score"],
        officer_id=1
    )

    try:
        # 1. GET /api/scans/{id}/export/csv
        resp_csv = client.get(f"/api/scans/{record.id}/export/csv")
        assert resp_csv.status_code == 200
        assert "text/csv" in resp_csv.headers["content-type"]
        assert f"filename=\"legal_metrology_scan_{record.id}.csv\"" in resp_csv.headers["content-disposition"]
        assert "=== LEGAL METROLOGY COMPLIANCE AUDIT RECORD ===" in resp_csv.text

        # 2. GET /api/scans/{id}/export/xlsx
        resp_xlsx = client.get(f"/api/scans/{record.id}/export/xlsx")
        assert resp_xlsx.status_code == 200
        assert "spreadsheetml.sheet" in resp_xlsx.headers["content-type"]
        assert f"filename=\"legal_metrology_scan_{record.id}.xlsx\"" in resp_xlsx.headers["content-disposition"]
        assert resp_xlsx.content.startswith(b"PK\x03\x04")

        # 3. GET /api/scans/{id}/export/docx
        resp_docx = client.get(f"/api/scans/{record.id}/export/docx")
        assert resp_docx.status_code == 200
        assert "wordprocessingml.document" in resp_docx.headers["content-type"]
        assert f"filename=\"show_cause_notice_scan_{record.id}.docx\"" in resp_docx.headers["content-disposition"]
        assert resp_docx.content.startswith(b"PK\x03\x04")

        # 4. GET /api/scans/export/xlsx (Bulk)
        resp_bulk = client.get("/api/scans/export/xlsx")
        assert resp_bulk.status_code == 200
        assert "spreadsheetml.sheet" in resp_bulk.headers["content-type"]
        assert "filename=\"legal_metrology_scans_bulk_" in resp_bulk.headers["content-disposition"]
        assert resp_bulk.content.startswith(b"PK\x03\x04")

        # 5. Non-existent scan returns 404
        resp_404_csv = client.get("/api/scans/99999/export/csv")
        assert resp_404_csv.status_code == 404

        resp_404_xlsx = client.get("/api/scans/99999/export/xlsx")
        assert resp_404_xlsx.status_code == 404

        resp_404_docx = client.get("/api/scans/99999/export/docx")
        assert resp_404_docx.status_code == 404

    finally:
        app.dependency_overrides.clear()


def test_font_size_disclaimer_in_exports(sample_scan_data):
    """Verifies that FONT_SIZE_DISCLAIMER is consistently present in Excel and DOCX exports."""
    from app.constants import FONT_SIZE_DISCLAIMER

    service = ExportService()

    # Excel Export Check
    xlsx_bytes = service.generate_scan_xlsx(sample_scan_data)
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
    assert "Legal Disclaimer" in wb.sheetnames
    disc_sheet = wb["Legal Disclaimer"]
    assert FONT_SIZE_DISCLAIMER in disc_sheet.cell(row=3, column=1).value

    # DOCX Show Cause Notice Check
    docx_bytes = service.generate_show_cause_docx(sample_scan_data, officer_name="Inspector R. Kumar")
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert FONT_SIZE_DISCLAIMER in full_text

